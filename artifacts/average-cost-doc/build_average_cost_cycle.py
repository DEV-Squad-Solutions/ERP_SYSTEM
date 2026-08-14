from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "C9D2DC"
GREEN = "1E6B4F"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(
    paragraph,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_keep(paragraph, with_next: bool = True, together: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell,
    top: int = 80,
    start: int = 120,
    bottom: int = 80,
    end: int = 120,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    tbl_pr = table._tbl.tblPr

    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    table_indent = tbl_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)


def set_paragraph_shading(paragraph, fill: str, border_color: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), border_color)
        borders.append(border)
    p_pr.append(borders)


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        r_pr = run._element.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        r_pr.append(rtl)
        set_run_font(run, name="Arial", size=10.5)


def add_numbering_definition(document: Document, decimal: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1000000 + abstract_id:08X}")
    abstract.append(nsid)
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{0xB1000000 + abstract_id:08X}")
    abstract.append(template)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    level.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if decimal else "•")
    level.append(lvl_text)

    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    first_num_position = next(
        (
            position
            for position, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_position, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)

    if "Formula" not in styles:
        formula = styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = styles["Formula"]
    formula.font.name = "Consolas"
    formula._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    formula.font.size = Pt(9.5)
    formula.font.color.rgb = RGBColor.from_string(NAVY)
    formula.paragraph_format.left_indent = Inches(0.2)
    formula.paragraph_format.right_indent = Inches(0.2)
    formula.paragraph_format.space_before = Pt(4)
    formula.paragraph_format.space_after = Pt(6)
    formula.paragraph_format.line_spacing = 1.15

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    note._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    note.font.size = Pt(9.5)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.15


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "MINIERP  /  INVENTORY COSTING REFERENCE"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line_spacing=1.0)
    set_run_font(paragraph.runs[0], size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Average Cost Cycle  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(paragraph)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_run = first_footer.paragraphs[0].add_run(
        "MiniErp technical reference  |  July 2026"
    )
    set_run_font(first_run, size=8.5, color=MUTED)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    set_keep(paragraph)


def add_body(
    document: Document,
    text: str,
    bold_lead: str | None = None,
) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    set_paragraph_spacing(paragraph)


def add_bullet(document: Document, text: str, bullet_num_id: int) -> None:
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, bullet_num_id)
    run = paragraph.add_run(text)
    set_run_font(run)
    set_paragraph_spacing(paragraph, after=4)


def add_numbered(document: Document, text: str, number_num_id: int) -> None:
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, number_num_id)
    run = paragraph.add_run(text)
    set_run_font(run)
    set_paragraph_spacing(paragraph, after=4)


def add_formula(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Formula")
    paragraph.add_run(text)
    set_paragraph_shading(paragraph, LIGHT_GRAY)


def add_callout(
    document: Document,
    label: str,
    text: str,
    fill: str = CALLOUT,
    label_color: str = NAVY,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=label_color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    set_paragraph_shading(paragraph, fill)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    header_fill: str = LIGHT_BLUE,
    font_size: float = 9.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_cells = table.rows[0].cells
    for index, text in enumerate(headers):
        cell = header_cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line_spacing=1.1)
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row_cells = table.add_row().cells
        for index, text in enumerate(row_data):
            cell = row_cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(text) < 25 and index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_spacing(paragraph, after=0, line_spacing=1.12)
            run = paragraph.add_run(text)
            set_run_font(run, size=font_size, color=INK)

    set_table_geometry(table, widths_dxa)
    document.add_paragraph(style="Small Note")


def add_arabic_error(document: Document, code: str, message: str) -> None:
    code_paragraph = document.add_paragraph()
    code_run = code_paragraph.add_run(code)
    set_run_font(code_run, name="Consolas", size=9.5, color=NAVY, bold=True)
    set_paragraph_spacing(code_paragraph, after=2)
    set_keep(code_paragraph)

    message_paragraph = document.add_paragraph()
    message_run = message_paragraph.add_run(message)
    set_run_font(message_run, name="Arial", size=10.5)
    set_paragraph_spacing(message_paragraph, after=7, line_spacing=1.2)
    set_rtl(message_paragraph)


def build_document(output_path: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_page(document)
    bullet_num_id = add_numbering_definition(document, decimal=False)
    number_num_id = add_numbering_definition(document, decimal=True)

    document.core_properties.title = (
        "MiniErp Perpetual Weighted-Average Inventory Costing Cycle"
    )
    document.core_properties.subject = (
        "Full business, data, transaction, replay, API, and migration reference"
    )
    document.core_properties.author = "MiniErp Engineering"
    document.core_properties.keywords = (
        "MiniErp, inventory, weighted average, average cost, negative stock, FIFO"
    )

    # Cover page: editorial_cover pattern with compact-reference typography.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(78)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("TECHNICAL REFERENCE GUIDE")
    set_run_font(run, size=10, color=BLUE, bold=True)
    set_paragraph_spacing(kicker, after=16, line_spacing=1.0)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Perpetual Weighted-Average\nInventory Costing Cycle")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Positive stock, negative-stock pending cost, FIFO revaluation, "
        "backdated replay, and current item/store balances"
    )

    context = document.add_paragraph()
    context.alignment = WD_ALIGN_PARAGRAPH.CENTER
    context_run = context.add_run("MiniErp  |  Final implementation reference")
    set_run_font(context_run, size=11, color=MUTED, bold=True)
    set_paragraph_spacing(context, before=18, after=4, line_spacing=1.0)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run("Version 1.0  |  29 July 2026")
    set_run_font(date_run, size=10, color=MUTED)
    set_paragraph_spacing(date, after=32, line_spacing=1.0)

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status.add_run(
        "Implementation complete - migration creation pending explicit approval"
    )
    set_run_font(status_run, size=10.5, color=GOLD, bold=True)
    set_paragraph_spacing(status, after=0, line_spacing=1.0)

    document.add_page_break()

    add_heading(document, "Contents", 1)
    contents = [
        "Purpose, scope, and invariants",
        "Partition key and persisted model",
        "Precision, rounding, and server ownership",
        "Company stock-check policy",
        "Complete movement-processing cycle",
        "Inbound and outbound cost-source rules",
        "Negative-stock pending-cost lifecycle",
        "Worked costing examples",
        "Backdated replay and stable movement identity",
        "Return linkage rules",
        "Atomicity, locking, and concurrency",
        "API, Swagger, seed, and frontend contracts",
        "Migration and historical-data backfill plan",
        "Verification matrix",
        "Operational checklist",
    ]
    for item in contents:
        add_numbered(document, item, number_num_id)
    add_bullet(document, "Appendix A. Arabic business messages", bullet_num_id)
    add_bullet(document, "Appendix B. Implementation reference", bullet_num_id)

    add_callout(
        document,
        "Reading rule",
        "Stock validation decides whether a quantity timeline is permitted. "
        "Inventory costing then assigns monetary cost to that permitted "
        "timeline. The two stages are related but must remain separate.",
    )

    add_heading(document, "1. Purpose, scope, and invariants", 1)
    add_body(
        document,
        "This document defines the complete perpetual weighted-average inventory "
        "cost cycle implemented by MiniErp. It covers movement source costs, "
        "positive and negative quantity behavior, FIFO coverage of pending "
        "outbound cost, historical replay, current balances, transactions, "
        "contracts, and migration expectations.",
    )

    for text in (
        "Cost is calculated independently for each CompanyId + StoreId + ItemId partition.",
        "Every active ItemMovement participates in chronological replay.",
        "Average cost is calculated only when quantity is positive.",
        "When QuantityAfter <= 0, AverageCostAfter and InventoryValueAfter are both zero.",
        "Pending outbound quantity is never finalized at zero cost.",
        "Purchase returns use the current weighted average, never the original purchase price.",
        "Matching movement identity and CreatedOn are preserved during document updates.",
        "Costing snapshots, allocations, and current balances are server-owned.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_callout(
        document,
        "Non-goals",
        "Inventory costing does not introduce document status, posting, "
        "cancellation, reversal, general-ledger entries, vouchers, or payment "
        "allocations.",
    )

    add_heading(document, "2. Partition key and persisted model", 1)
    add_formula(document, "InventoryCostingKey = CompanyId + StoreId + ItemId")
    add_body(
        document,
        "All quantity, average-cost, inventory-value, pending-cost, allocation, "
        "locking, and replay decisions are isolated by this tenant-safe key.",
    )

    add_heading(document, "2.1 Persisted entities", 2)
    add_table(
        document,
        ["Entity", "Purpose", "Important persisted fields"],
        [
            [
                "ItemMovement",
                "Canonical chronological quantity and cost event.",
                "UnitCost; TotalCost; QuantityAfter; AverageCostAfter; "
                "InventoryValueAfter; PendingCostQuantity; CostStatus.",
            ],
            [
                "InventoryCostAllocation",
                "Derived audit row linking future inbound cost to an earlier pending outbound.",
                "CompanyId; StoreId; ItemId; OutboundMovementId; "
                "InboundMovementId; Quantity; UnitCost; TotalCost; CreatedOn.",
            ],
            [
                "ItemStoreBalance",
                "Current server-maintained projection for one item/store.",
                "Quantity; AverageCost; InventoryValue; RowVersion.",
            ],
            [
                "InvoiceLine",
                "Invoice input and optional linked sales-return source.",
                "Price; SourceInvoiceLineId; ReturnUnitCost.",
            ],
            [
                "StockAdjustmentLine",
                "Manual adjustment input.",
                "UnitCost is required only for an Increase.",
            ],
        ],
        [2200, 3000, 4160],
    )

    add_heading(document, "2.2 Keys, indexes, and constraints", 2)
    add_table(
        document,
        ["Object", "EF Core / database design", "Reason"],
        [
            [
                "ItemMovement",
                "Primary key Id; alternate key (CompanyId, Id).",
                "Supports tenant-safe composite allocation foreign keys.",
            ],
            [
                "ItemMovement timeline",
                "Filtered index on CompanyId, StoreId, ItemId, MovementDate, CreatedOn, Id.",
                "Deterministic active replay.",
            ],
            [
                "Pending movements",
                "Filtered index includes CostStatus and chronological columns.",
                "Locates Pending and PartiallyCosted movements.",
            ],
            [
                "InventoryCostAllocation",
                "Unique CompanyId, OutboundMovementId, InboundMovementId.",
                "One rebuilt allocation per inbound/outbound pair.",
            ],
            [
                "ItemStoreBalance",
                "Composite primary key CompanyId, StoreId, ItemId; RowVersion.",
                "One concurrent current balance per costing partition.",
            ],
            [
                "InvoiceLine source",
                "Tenant-safe optional self-reference by CompanyId and SourceInvoiceLineId.",
                "Prevents cross-company return-cost linkage.",
            ],
        ],
        [2200, 3800, 3360],
    )

    add_body(
        document,
        "Database checks enforce nonnegative monetary cost, pending quantity not "
        "exceeding QuantityOut, exactly one movement direction, and zero "
        "average/value whenever the running quantity is nonpositive.",
    )

    add_heading(document, "3. Precision, rounding, and server ownership", 1)
    add_table(
        document,
        ["Value family", "SQL precision", "Rounding / ownership"],
        [
            ["Quantity", "decimal(18,6)", "AwayFromZero; input quantity is validated."],
            [
                "Unit and average cost",
                "decimal(24,8)",
                "AwayFromZero; source input only where explicitly allowed.",
            ],
            [
                "Total cost and inventory value",
                "decimal(28,8)",
                "AwayFromZero; always server-calculated.",
            ],
        ],
        [2600, 2300, 4460],
    )

    add_heading(document, "3.1 Client-entered cost fields", 2)
    for text in (
        "Purchase invoice line Price.",
        "Opening-balance line Price, presented to the user as Unit Cost.",
        "Stock-adjustment Increase UnitCost.",
        "Unlinked Sales Return ReturnUnitCost only when no positive current average exists.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "3.2 Server-calculated fields", 2)
    for text in (
        "ItemMovement CostStatus, PendingCostQuantity, UnitCost for outbound movements, TotalCost, QuantityAfter, AverageCostAfter, and InventoryValueAfter.",
        "Every InventoryCostAllocation field.",
        "ItemStoreBalance Quantity, AverageCost, InventoryValue, and RowVersion.",
        "All cost snapshots returned by invoice, adjustment, and opening-balance APIs.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_callout(
        document,
        "Important",
        "A Stock Adjustment Decrease must omit UnitCost. The server uses the "
        "weighted average immediately before that outbound movement.",
        fill="FFF8E8",
        label_color=GOLD,
    )

    add_heading(document, "4. Company stock-check policy", 1)
    add_body(
        document,
        "Each company selects one StockBalanceCheckMode. This controls whether "
        "the quantity timeline may become negative; it does not disable costing.",
    )
    add_table(
        document,
        ["Mode", "Quantity validation", "Costing behavior"],
        [
            ["None", "No negative-stock rejection.", "Replay and pending-cost logic still run."],
            [
                "DateCheck",
                "Reject any negative point in the chronological timeline.",
                "Replay runs only after validation succeeds.",
            ],
            [
                "FinalCheck",
                "Reject only a negative resulting final balance.",
                "Earlier negative points may be costed through pending FIFO.",
            ],
            [
                "Both",
                "Apply DateCheck and FinalCheck.",
                "Replay runs only after both checks succeed.",
            ],
        ],
        [1600, 3900, 3860],
    )

    add_body(
        document,
        "Outbound create, update, and delete run the configured check. Inbound "
        "update and delete also run it because they may remove quantity that "
        "supports later outbound movements. A new inbound create skips stock "
        "validation because it only adds quantity, but it still triggers costing replay.",
    )

    add_heading(document, "5. Complete movement-processing cycle", 1)
    steps = [
        "Start one Serializable database transaction for the document operation.",
        "Validate request shape, tenant ownership, active product store, active item and unit, and document header RowVersion where applicable.",
        "Derive every old and new CompanyId + StoreId + ItemId key affected by the change.",
        "Lock balance keys in deterministic CompanyId, StoreId, ItemId order. SQL Server uses UPDLOCK and HOLDLOCK, including a range lock when no balance row exists.",
        "Run the company stock-check policy against the exact proposed quantity timeline.",
        "Reconcile document movements while preserving matching ItemMovement.Id and CreatedOn.",
        "Load the full active movement timeline for each affected key ordered by MovementDate, CreatedOn, Id.",
        "Physically remove the affected derived InventoryCostAllocation rows.",
        "Replay every inbound and outbound movement, rebuild pending FIFO allocations, and refresh all movement snapshots.",
        "Apply the last replay state to ItemStoreBalance.",
        "Save the document, movements, snapshots, allocations, and balances and commit. Any failure rolls everything back.",
    ]
    cycle_num_id = add_numbering_definition(document, decimal=True)
    for step in steps:
        add_numbered(document, step, cycle_num_id)

    add_callout(
        document,
        "Separation of concerns",
        "Stock validation evaluates quantity permission. Costing replay evaluates "
        "monetary value. Validation must finish before final costing is saved.",
    )

    add_heading(document, "6. Inbound and outbound cost-source rules", 1)
    add_heading(document, "6.1 Inbound movements", 2)
    add_table(
        document,
        ["Movement", "Unit-cost source", "Special rule"],
        [
            ["Purchase", "Purchase invoice line Price.", "Recalculates weighted average."],
            [
                "Opening balance",
                "Opening-balance line Price.",
                "Creates an OpeningBalance ItemMovement.",
            ],
            [
                "Adjustment Increase",
                "Required StockAdjustmentLine.UnitCost.",
                "Client cost must be nonnegative.",
            ],
            [
                "Linked Sales Return",
                "Final UnitCost of the original Sales movement.",
                "Source sale must precede the return and be fully costed.",
            ],
            [
                "Unlinked Sales Return",
                "Current positive average; otherwise ReturnUnitCost.",
                "Fallback is required only without a positive average.",
            ],
            [
                "Transfer In",
                "Stored source movement cost.",
                "Rejected when no source cost is available.",
            ],
        ],
        [2100, 3400, 3860],
    )

    add_heading(document, "6.2 Outbound movements", 2)
    for text in (
        "Sales use the current weighted average immediately before the movement.",
        "Stock Adjustment Decrease uses the current weighted average immediately before the movement.",
        "Purchase Return uses the current weighted average immediately before the movement and never the original purchase price.",
        "An outbound movement does not change a positive average unless remaining quantity becomes zero or negative.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "6.3 Positive-stock formulas", 2)
    add_formula(
        document,
        "InboundTotalCost      = QuantityIn * SourceUnitCost\n"
        "QuantityAfter        = PreviousQuantity + QuantityIn\n"
        "InventoryValueAfter  = PreviousInventoryValue + InboundTotalCost\n"
        "AverageCostAfter     = InventoryValueAfter / QuantityAfter\n"
        "Condition: calculate average only when QuantityAfter > 0",
    )
    add_formula(
        document,
        "OutboundUnitCost     = PreviousAverageCost\n"
        "OutboundTotalCost    = QuantityOut * OutboundUnitCost\n"
        "QuantityAfter        = PreviousQuantity - QuantityOut\n"
        "InventoryValueAfter  = PreviousInventoryValue - OutboundTotalCost",
    )

    add_heading(document, "7. Negative-stock pending-cost lifecycle", 1)
    add_body(
        document,
        "When policy permits an outbound quantity greater than available positive "
        "stock, the movement is logically divided into covered and pending portions.",
    )
    add_formula(
        document,
        "CoveredQuantity      = min(QuantityOut, max(PreviousQuantity, 0))\n"
        "PendingCostQuantity  = QuantityOut - CoveredQuantity\n"
        "CoveredCost          = CoveredQuantity * PreviousAverageCost",
    )

    add_heading(document, "7.1 Cost status", 2)
    add_table(
        document,
        ["Status", "Meaning", "UnitCost rule"],
        [
            ["Final", "Completely costed without future dependency.", "Populated."],
            [
                "PartiallyCosted",
                "Some outbound quantity is costed and some remains pending.",
                "Null until fully covered.",
            ],
            [
                "Pending",
                "The full outbound quantity awaits future inbound cost.",
                "Null.",
            ],
            [
                "Revalued",
                "A previously pending movement is now fully covered.",
                "TotalCost / QuantityOut.",
            ],
        ],
        [1800, 4700, 2860],
    )

    add_heading(document, "7.2 FIFO coverage by later inbound", 2)
    fifo_num_id = add_numbering_definition(document, decimal=True)
    for text in (
        "Pending outbound movements are queued by MovementDate, CreatedOn, Id.",
        "A future inbound first covers the earliest pending outbound quantity.",
        "Each cover writes one InventoryCostAllocation with quantity, source unit cost, and total cost.",
        "Only the remaining inbound quantity enters positive inventory and participates in the new average.",
        "When the pending quantity reaches zero, the outbound becomes Revalued and receives its final effective UnitCost.",
    ):
        add_numbered(document, text, fifo_num_id)

    add_callout(
        document,
        "Hard invariant",
        "When QuantityAfter <= 0, AverageCostAfter = 0 and "
        "InventoryValueAfter = 0. Never divide inventory value by a zero or "
        "negative quantity.",
        fill="FFF2F2",
        label_color=RED,
    )

    document.add_page_break()
    add_heading(document, "8. Worked costing examples", 1)
    add_heading(document, "8.1 Positive weighted-average cycle", 2)
    add_body(
        document,
        "Start with 10 units at cost 5.00. Purchase 10 units at 7.00, sell "
        "8 units, increase by 3 units at 9.00, then return 5 units to the supplier.",
    )
    add_table(
        document,
        ["Movement", "Calculation", "Snapshot after movement"],
        [
            ["Opening 10 @ 5", "Value = 50", "Qty 10; Avg 5; Value 50"],
            [
                "Purchase 10 @ 7",
                "New value = 50 + 70",
                "Qty 20; Avg 6; Value 120",
            ],
            [
                "Sale 8",
                "COGS = 8 * 6 = 48",
                "Qty 12; Avg 6; Value 72",
            ],
            [
                "Adjustment Increase 3 @ 9",
                "New value = 72 + 27",
                "Qty 15; Avg 6.6; Value 99",
            ],
            [
                "Purchase Return 5",
                "Cost = 5 * 6.6 = 33",
                "Qty 10; Avg 6.6; Value 66",
            ],
        ],
        [2500, 3300, 3560],
    )

    add_heading(document, "8.2 Partial negative stock and revaluation", 2)
    add_body(
        document,
        "Previous quantity is 3 at average cost 10. An outbound movement issues "
        "5 units. A later purchase brings 10 units at cost 12.",
    )
    add_formula(
        document,
        "Outbound: covered 3 * 10 = 30; pending 2; QuantityAfter = -2\n"
        "Snapshot: TotalCost = 30; UnitCost = null; CostStatus = PartiallyCosted\n"
        "Inbound: allocate 2 * 12 = 24; remaining inbound = 8\n"
        "Revalued outbound: TotalCost = 54; UnitCost = 54 / 5 = 10.8\n"
        "Final balance: Quantity = 8; AverageCost = 12; InventoryValue = 96",
    )

    add_heading(document, "8.3 Multiple inbound covers", 2)
    add_body(
        document,
        "An outbound of 6 units occurs from zero stock. Later inbound movements "
        "arrive as 2 @ 10, 3 @ 12, and 4 @ 15.",
    )
    add_formula(
        document,
        "After outbound: Pending = 6; TotalCost = 0\n"
        "Inbound 2 @ 10: allocate 20; Pending = 4\n"
        "Inbound 3 @ 12: allocate 36; Pending = 1; accumulated cost = 56\n"
        "Inbound 4 @ 15: allocate 15; outbound TotalCost = 71\n"
        "Final outbound UnitCost = 71 / 6 = 11.83333333\n"
        "Remaining positive stock = 3 @ 15; InventoryValue = 45",
    )

    add_heading(document, "8.4 Accounting reconciliation", 2)
    add_formula(
        document,
        "Opening inventory value + total inbound value\n"
        "= final inventory value + total fully costed outbound cost",
    )
    add_body(
        document,
        "The equation closes when all outbound movements are fully costed. If "
        "Pending or PartiallyCosted movements remain, the unresolved quantity "
        "is reported explicitly and the final cost equation is not yet closed.",
    )

    add_heading(document, "9. Backdated replay and stable movement identity", 1)
    add_heading(document, "9.1 Replay triggers", 2)
    replay_triggers = (
        "create; edit; delete; restore; movement-date change; quantity change; "
        "cost change; item change; store change; company change; direction "
        "change; movement-type change; source-linkage change; soft-delete change"
    )
    add_body(document, f"Full replay is required after any of: {replay_triggers}.")
    add_body(
        document,
        "Every old and new affected partition is replayed. Replay includes all "
        "subsequent inbound and outbound movements, not only inbound movements. "
        "When the safe dependency boundary is uncertain, MiniErp replays the "
        "full active timeline for the item/store.",
    )

    add_heading(document, "9.2 Stable movement identity", 2)
    for text in (
        "Update matching movements in place.",
        "Preserve matching ItemMovement.Id and CreatedOn.",
        "Soft-delete movements for removed lines.",
        "Create a new movement only for a newly added line.",
        "Never delete and recreate all document movements during a normal update.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "9.3 Allocation rebuild", 2)
    add_body(
        document,
        "InventoryCostAllocation is server-derived and rebuildable. Replay "
        "physically deletes the affected allocation rows, then rebuilds them "
        "deterministically. The unique movement-pair constraint prevents "
        "ambiguous duplicate allocations.",
    )

    add_heading(document, "10. Return linkage rules", 1)
    add_heading(document, "10.1 Linked Sales Return", 2)
    for text in (
        "SourceInvoiceLineId must identify an active Sales line in the same company, store, and item.",
        "The source Sales movement must precede the return in deterministic movement order.",
        "The source movement must be Final or Revalued and have a non-null UnitCost.",
        "Pending and PartiallyCosted source sales are rejected.",
        "An active linked sales return blocks soft deletion of its source Sales invoice.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "10.2 Unlinked Sales Return", 2)
    add_body(
        document,
        "Use the current positive average when available. Require "
        "ReturnUnitCost only when no positive average exists. ReturnUnitCost is "
        "a fallback source input, not a server snapshot field.",
    )

    add_heading(document, "10.3 Purchase Return", 2)
    add_callout(
        document,
        "Approved policy",
        "A purchase return is a normal outbound movement. It uses the current "
        "weighted average immediately before the movement and never uses the "
        "original purchase price.",
        fill="EEF8F3",
        label_color=GREEN,
    )

    add_heading(document, "11. Atomicity, locking, and concurrency", 1)
    add_heading(document, "11.1 Transaction boundary", 2)
    add_body(
        document,
        "Invoices, stock adjustments, opening balances, and inventory-count "
        "reconciliation execute their movement-producing workflows inside one "
        "Serializable SQL Server transaction. A failure in validation, document "
        "concurrency, costing, allocation rebuilding, or balance persistence "
        "rolls the complete operation back.",
    )

    add_heading(document, "11.2 Deterministic balance locking", 2)
    add_body(
        document,
        "Affected keys are ordered by CompanyId, StoreId, ItemId. Existing "
        "ItemStoreBalance rows are locked with UPDLOCK and HOLDLOCK. An exact "
        "key lookup under Serializable isolation range-locks a missing key so "
        "two writers cannot independently create the same balance.",
    )

    add_heading(document, "11.3 Concurrency tokens", 2)
    for text in (
        "Document aggregates retain their existing header RowVersion contract.",
        "The client must send the originally returned header token; EF Core uses that exact token as OriginalValue.",
        "Line-only aggregate changes touch the header so its RowVersion advances.",
        "ItemStoreBalance has its own SQL Server RowVersion for current-balance write conflicts.",
        "DbUpdateConcurrencyException rolls back the transaction and returns the feature's Arabic reload-and-retry conflict.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "12. API, Swagger, seed, and frontend contracts", 1)
    add_heading(document, "12.1 API response fields", 2)
    add_body(
        document,
        "Invoice, stock-adjustment, and opening-balance line responses expose "
        "movement CostStatus, PendingCostQuantity where applicable, UnitCost, "
        "TotalCost, QuantityAfter, AverageCostAfter, and InventoryValueAfter. "
        "Invoice item-balance responses expose current Balance, AverageCost, "
        "and InventoryValue.",
    )

    add_heading(document, "12.2 Request rules", 2)
    for text in (
        "Stock Adjustment Increase: UnitCost is required for every line.",
        "Stock Adjustment Decrease: UnitCost must be omitted.",
        "Inventory Count reconciliation: IncreaseCosts must contain exactly one UnitCost for each positive difference and no extras.",
        "Sales Return: optional SourceInvoiceLineId; otherwise ReturnUnitCost only when no positive average exists.",
        "Server-calculated costing fields are never accepted from the client.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "12.3 Frontend behavior", 2)
    for text in (
        "Company screen keeps company information and stock options in separate tabs; the stock-check dropdown has exactly four modes.",
        "Every paginated Stock Adjustment item includes the complete deterministic line collection.",
        "Adjustment, invoice, and opening-balance screens show unit cost, average cost, and inventory value.",
        "Inventory Count requests a unit cost for each positive adjustment before reconciliation.",
        "Sales Return supports source-line linkage and fallback ReturnUnitCost.",
    ):
        add_bullet(document, text, bullet_num_id)

    add_heading(document, "12.4 Seed behavior", 2)
    add_body(
        document,
        "Development seed ensures opening-balance movements exist, updates "
        "matching seeded movements in place, then runs full costing replay for "
        "each seeded company to populate movement snapshots, allocations, and "
        "current balances idempotently.",
    )

    add_heading(document, "13. Migration and historical-data backfill plan", 1)
    add_callout(
        document,
        "Current status",
        "The entity and EF Core configuration are implemented. No perpetual "
        "costing migration has been created or applied. Migration generation "
        "requires a separate explicit approval.",
        fill="FFF8E8",
        label_color=GOLD,
    )

    migration_steps = [
        "Add nullable UnitCost and required snapshot/status fields to ItemMovements with safe temporary defaults for existing rows.",
        "Create InventoryCostAllocations with tenant-safe composite foreign keys and the unique movement-pair constraint.",
        "Create ItemStoreBalances with the composite primary key, monetary checks, and SQL Server RowVersion.",
        "Add SourceInvoiceLineId and ReturnUnitCost to InvoiceLines and UnitCost to StockAdjustmentLines.",
        "Backfill one stable OpeningBalance ItemMovement for every active opening-balance line that has no matching movement.",
        "Resolve source cost for every historical inbound movement from its document line.",
        "Replay every active CompanyId + StoreId + ItemId timeline in deterministic order.",
        "Populate all movement snapshots, rebuild pending allocations, and create current ItemStoreBalance rows.",
        "Validate nonnegative-cost and nonpositive-state constraints before enabling them against historical data.",
        "Review Up, Down, generated snapshot, indexes, foreign keys, and existing-data SQL before applying the migration.",
    ]
    migration_num_id = add_numbering_definition(document, decimal=True)
    for step in migration_steps:
        add_numbered(document, step, migration_num_id)

    add_heading(document, "14. Verification matrix", 1)
    add_table(
        document,
        ["Area", "Required scenarios"],
        [
            [
                "Positive costing",
                "Weighted inbound average; outbound current average; quantity "
                "to zero; average/value reset.",
            ],
            [
                "Negative costing",
                "Fully pending; partially costed; one/multiple inbound covers; "
                "one inbound covers multiple outbounds; exact zero and positive recovery.",
            ],
            [
                "Replay",
                "Backdated create; quantity/cost/date edit; item/store move; "
                "delete; restore; full later inbound/outbound recalculation.",
            ],
            [
                "Identity",
                "Matching updates preserve movement Id and CreatedOn; removed "
                "lines soft-delete movements.",
            ],
            [
                "Returns",
                "Purchase return current average; linked sale final/revalued; "
                "pending-source rejection; unlinked fallback.",
            ],
            [
                "Allocations",
                "Deterministic rebuild; unique pair; FIFO chronology; correct "
                "effective outbound UnitCost.",
            ],
            [
                "Balance",
                "Current balance equals final movement snapshot; quantity/value "
                "and accounting equation reconcile.",
            ],
            [
                "Cross-cutting",
                "Tenant isolation; active store/item validation; transaction "
                "rollback; stale RowVersion; Arabic errors; Swagger; frontend build.",
            ],
        ],
        [2500, 6860],
    )

    add_body(
        document,
        "Verification recorded when this guide was prepared: 499 backend tests "
        "passed, .NET formatting verification passed, and the frontend "
        "production build passed.",
    )

    add_heading(document, "15. Operational checklist", 1)
    checklist = (
        "Confirm the document transaction uses Serializable isolation.",
        "Derive both old and new item/store keys before changing movements.",
        "Acquire balance locks in deterministic order before stock validation and movement writes.",
        "Run the company's selected stock-check mode.",
        "Preserve matching movement identity and CreatedOn.",
        "Replay MovementDate, CreatedOn, Id across all later active movements.",
        "Rebuild affected allocation rows deterministically.",
        "Confirm nonpositive quantity has zero average and inventory value.",
        "Confirm pending outbound UnitCost remains null.",
        "Confirm successful responses return refreshed snapshots and RowVersion.",
        "Confirm rollback leaves no partial document, movement, allocation, or balance state.",
        "Do not create or apply a migration until the generated design is approved.",
    )
    for item in checklist:
        add_bullet(document, f"Check: {item}", bullet_num_id)

    add_heading(document, "Appendix A. Arabic business messages", 1)
    add_body(
        document,
        "The following messages are part of the implemented business contract.",
    )
    add_arabic_error(
        document,
        "Inventory.SalesReturnSourceCostPending",
        "لا يمكن احتساب تكلفة مرتجع البيع لأن حركة البيع الأصلية لم تكتمل تكلفتها بعد.",
    )
    add_arabic_error(
        document,
        "Inventory.ReturnUnitCostRequired",
        "يجب إدخال تكلفة وحدة مرتجع البيع عند عدم توفر متوسط تكلفة موجب.",
    )
    add_arabic_error(
        document,
        "StockAdjustments.UnitCostRequired",
        "يجب إدخال تكلفة الوحدة لكل سطر عند زيادة المخزون.",
    )
    add_arabic_error(
        document,
        "StockAdjustments.UnitCostNotAllowed",
        "لا يجوز إدخال تكلفة الوحدة في تسوية الخصم؛ يستخدم الخادم متوسط التكلفة الحالي.",
    )
    add_arabic_error(
        document,
        "Inventory.MovementCostSourceMissing",
        "تعذر العثور على مصدر تكلفة حركة المخزون.",
    )

    add_heading(document, "Appendix B. Implementation reference", 1)
    references = (
        "MiniErp.Domain/Entities/Inventory/ItemMovement.cs",
        "MiniErp.Domain/Entities/Inventory/InventoryCostAllocation.cs",
        "MiniErp.Domain/Entities/Inventory/ItemStoreBalance.cs",
        "MiniErp.Infrastructure/Services/Inventory/InventoryCostingService.cs",
        "MiniErp.Infrastructure/Persistence/Configurations/ItemMovementConfiguration.cs",
        "MiniErp.Infrastructure/Persistence/Configurations/InventoryCostAllocationConfiguration.cs",
        "MiniErp.Infrastructure/Persistence/Configurations/ItemStoreBalanceConfiguration.cs",
        "MiniErp.Infrastructure/Services/Invoices/InvoiceService*.cs",
        "MiniErp.Infrastructure/Services/StockAdjustments/StockAdjustmentService.cs",
        "MiniErp.Infrastructure/Services/StockOpeningBalances/StockOpeningBalanceService.cs",
        "MiniErp.Infrastructure/Services/InventoryCounts/InventoryCountService.cs",
        "FEATURE_DEVELOPMENT_GUIDE.md - Perpetual weighted-average inventory costing",
    )
    for reference in references:
        add_bullet(document, reference, bullet_num_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_document(args.output)


if __name__ == "__main__":
    main()
